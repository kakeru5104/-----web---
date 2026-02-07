from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_specification_docx():
    # ドキュメントの作成
    doc = Document()

    # --- タイトル ---
    title = doc.add_heading('推しメンクイズ機能 仕様書', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- メタ情報 ---
    p = doc.add_paragraph()
    p.add_run('プロジェクト名: ').bold = True
    p.add_run('GRADUATION LIVE \'26 公式サイト\n')
    p.add_run('作成日: ').bold = True
    p.add_run('2026年2月4日\n')
    p.add_run('バージョン: ').bold = True
    p.add_run('1.0')

    # --- 1. 概要 ---
    doc.add_heading('1. 概要', level=1)
    doc.add_paragraph(
        '本機能は、Webサイト訪問者が卒業メンバーに関する知識を試すことができる「推しメンクイズ」コンテンツである。'
        '特定日時までのロック機能、ニックネーム登録、クイズ回答、結果発表、およびデータベースへのスコア保存機能を有する。'
    )

    # --- 2. 画面遷移と機能詳細 ---
    doc.add_heading('2. 画面遷移と機能詳細', level=1)

    # 関数: 各セクションを追加するヘルパー
    def add_section(heading, img_files, text_content):
        doc.add_heading(heading, level=2)
        
        # 画像プレースホルダー
        p_img = doc.add_paragraph()
        run = p_img.add_run(f'【ここに {img_files} を貼り付けてください】')
        run.font.color.rgb = RGBColor(255, 0, 0) # 赤文字
        run.bold = True
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # テキスト内容
        for line in text_content:
            doc.add_paragraph(line, style='List Bullet')

    # 2-1
    add_section('2-1. コンテンツ解禁前（ロック画面）', '1.png', [
        '概要: キャンペーン開始日時より前にアクセスした場合の表示。',
        '表示条件: 現在時刻 < 2026年3月21日 19:00',
        'UI要素: タイトル「SECRET GAME」、ステータス「COMING SOON...」、解禁日時「2026.03.21 19:00 UNLOCK」',
        '機能: ユーザー操作不可。指定日時を過ぎると自動（またはリロード）でスタート画面へ切り替わる。'
    ])

    # 2-2
    add_section('2-2. スタート画面（解禁後）', '2.png, 3.png', [
        '概要: クイズの参加登録画面。',
        '表示条件: 現在時刻 ≧ 2026年3月21日 19:00',
        'UI要素: タイトル「推しメンクイズ 全問正解チャレンジ」、ニックネーム入力フォーム、STARTボタン',
        '機能: ニックネーム未入力時は遷移不可。',
        '状態遷移: 画像3.pngのように名前を入力しSTARTを押すと次へ遷移。'
    ])

    # 2-3
    add_section('2-3. メンバー選択画面', '4.png', [
        '概要: クイズに挑戦したい対象メンバーを選択する画面。',
        'UI要素: タイトル「SELECT MEMBER」、メンバーリスト、BACKボタン',
        '機能: リストタップで該当メンバーのクイズデータを読み込み遷移。縦スクロール可能。'
    ])

    # 2-4
    add_section('2-4. クイズ出題画面', '5.png, 6.png', [
        '概要: 4択クイズを出題・回答する画面。',
        'UI要素: 問題番号、問題文、選択肢ボタン、決定ボタン',
        '機能(初期): 画像5.pngの状態。',
        '機能(選択時): 画像6.pngの状態。選択肢をタップすると反転し、決定ボタンで次へ進む。全8問。'
    ])

    # 2-5
    add_section('2-5. 結果画面（通常）', '7.png', [
        '概要: 全問正解できなかった場合の結果表示（正解数 0〜7問）。',
        'UI要素: スコア（例: 3/8）、メッセージ（例: 出直してこい！）、RETRYボタン',
        '機能: DBへプレイ履歴（スコア、名前）を保存する。RETRYで最初に戻る。'
    ])

    # 2-6
    add_section('2-6. 結果画面（完全制覇）', '8.png', [
        '概要: 全問正解した場合の特別な結果表示（正解数 8問）。',
        'UI要素: 金色の背景、PERFECT!!の文字、ユーザー名の強調表示。',
        '機能: 通常同様にDB保存を行う。'
    ])

    # --- 3. データベース仕様 ---
    doc.add_heading('3. データベース仕様', level=1)
    
    # 画像プレースホルダー
    p_img = doc.add_paragraph()
    run = p_img.add_run('【ここに 9.png を貼り付けてください】')
    run.font.color.rgb = RGBColor(255, 0, 0)
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('テーブル名: game_results')

    # 表の作成
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # ヘッダー
    hdr_cells = table.rows[0].cells
    headers = ['カラム名', 'データ型', '説明', 'サンプル値']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True

    # データ行
    data = [
        ('id', 'int8', '主キー（自動採番）', '7, 8'),
        ('created_at', 'timestamptz', 'データ作成日時', '2026-02-04 09:34...'),
        ('player_name', 'text', '入力したニックネーム', 'test_user'),
        ('target_member', 'text', 'クイズ対象メンバー', '田中 太郎'),
        ('score', 'int2', '獲得スコア（0〜8）', '3, 8')
    ]

    for item in data:
        row_cells = table.add_row().cells
        for i, val in enumerate(item):
            row_cells[i].text = val

    doc.add_paragraph('\n保存タイミング: 結果画面が表示された時点。\n用途: プレイ履歴の集計、管理。')

    # 保存
    file_name = 'pushimen_quiz_spec.docx'
    doc.save(file_name)
    print(f'{file_name} を作成しました。')

if __name__ == '__main__':
    create_specification_docx()